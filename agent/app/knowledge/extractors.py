"""Text extraction for architecture knowledge source files."""

from __future__ import annotations

from pathlib import Path
from typing import Iterable

from agent.app.knowledge.models import KnowledgeDocument

SUPPORTED_EXTENSIONS = {".md", ".txt", ".pdf", ".docx"}


def source_type_for_path(path: Path) -> str:
    """Map file extensions to coarse source types used in citations."""
    if path.suffix.lower() == ".pdf":
        return "book"
    if path.suffix.lower() == ".docx":
        return "docx"
    if path.suffix.lower() == ".md":
        return "markdown"
    return "text"


def supported_files(path: Path) -> Iterable[Path]:
    """Yield supported files from a file or directory path in stable order."""
    if path.is_file():
        if path.suffix.lower() in SUPPORTED_EXTENSIONS:
            yield path
        return

    for candidate in sorted(path.rglob("*")):
        if candidate.is_file() and candidate.suffix.lower() in SUPPORTED_EXTENSIONS:
            yield candidate


def extract_document(path: str | Path) -> KnowledgeDocument:
    """Extract text from markdown, text, or PDF files."""
    source = Path(path)
    suffix = source.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported knowledge file extension: {source.suffix}")

    if suffix == ".pdf":
        text = _extract_pdf_text(source)
    elif suffix == ".docx":
        text = _extract_docx_text(source)
    else:
        text = source.read_text(encoding="utf-8")

    title = source.stem.replace("_", " ").replace("-", " ").strip() or source.name
    return KnowledgeDocument(
        title=title,
        source_type=source_type_for_path(source),
        path=str(source),
        text=text,
        metadata={"filename": source.name},
    )


def _extract_pdf_text(path: Path) -> str:
    try:
        from pypdf import PdfReader  # type: ignore[reportMissingImports]
    except Exception as exc:  # pragma: no cover - dependency failure path
        raise RuntimeError("PDF ingestion requires the pypdf package.") from exc

    reader = PdfReader(str(path))
    pages: list[str] = []
    for page_number, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        if page_text.strip():
            pages.append(f"\n\n[page {page_number}]\n{page_text.strip()}")
    return "\n".join(pages).strip()


def _extract_docx_text(path: Path) -> str:
    try:
        from docx import Document  # type: ignore[reportMissingImports]
    except Exception as exc:  # pragma: no cover - dependency failure path
        raise RuntimeError("DOCX ingestion requires the python-docx package.") from exc

    document = Document(str(path))
    return "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()).strip()
